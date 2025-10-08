"""
Document ingestion pipeline with OpenAI embeddings and vector storage
"""
import os
import hashlib
from typing import List, Dict, Any, Optional
from pathlib import Path
import asyncio
import aiofiles
from datetime import datetime

# Document processing
import PyPDF2
from docx import Document as DocxDocument
import pandas as pd

# LangChain and embeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LangChainDocument

# Open source embeddings
from .open_source_embeddings import get_embedding_service

# Vector stores
from langchain_community.vectorstores import Chroma
import chromadb
from chromadb.config import Settings as ChromaSettings

# LlamaIndex
from llama_index.core import Document as LlamaDocument
from llama_index.core.node_parser import SimpleNodeParser
from llama_index.core.extractors import TitleExtractor, QuestionsAnsweredExtractor

from ..core.config import settings
from ..models.document import DocumentChunk


class DocumentProcessor:
    """Handles document processing and text extraction"""
    
    @staticmethod
    async def extract_text_from_file(file_path: str, file_type: str) -> str:
        """Extract text content from various file types"""
        try:
            if file_type.lower() == '.pdf':
                return await DocumentProcessor._extract_from_pdf(file_path)
            elif file_type.lower() == '.docx':
                return await DocumentProcessor._extract_from_docx(file_path)
            elif file_type.lower() in ['.txt', '.md']:
                return await DocumentProcessor._extract_from_text(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
        except Exception as e:
            raise Exception(f"Error extracting text from {file_path}: {str(e)}")
    
    @staticmethod
    async def _extract_from_pdf(file_path: str) -> str:
        """Extract text from PDF file"""
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        return text.strip()
    
    @staticmethod
    async def _extract_from_docx(file_path: str) -> str:
        """Extract text from DOCX file"""
        doc = DocxDocument(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text.strip()
    
    @staticmethod
    async def _extract_from_text(file_path: str) -> str:
        """Extract text from plain text files"""
        async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
            return await file.read()


class TextChunker:
    """Handles text chunking for optimal embedding and retrieval"""
    
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.CHUNK_SIZE,
            chunk_overlap=settings.CHUNK_OVERLAP,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
        
        # LlamaIndex node parser for advanced chunking
        self.node_parser = SimpleNodeParser.from_defaults(
            chunk_size=settings.CHUNK_SIZE,
            chunk_overlap=settings.CHUNK_OVERLAP
        )
    
    def chunk_text(self, text: str, metadata: Dict[str, Any] = None) -> List[DocumentChunk]:
        """Split text into chunks with metadata"""
        if not metadata:
            metadata = {}
        
        # Use LangChain text splitter
        chunks = self.text_splitter.split_text(text)
        
        document_chunks = []
        for i, chunk in enumerate(chunks):
            chunk_metadata = {
                **metadata,
                "chunk_index": i,
                "chunk_size": len(chunk),
                "total_chunks": len(chunks)
            }
            
            document_chunks.append(DocumentChunk(
                document_id=metadata.get("document_id", 0),
                chunk_index=i,
                content=chunk,
                metadata=chunk_metadata
            ))
        
        return document_chunks
    
    def chunk_with_llamaindex(self, text: str, metadata: Dict[str, Any] = None) -> List[DocumentChunk]:
        """Advanced chunking using LlamaIndex"""
        if not metadata:
            metadata = {}
        
        # Create LlamaIndex document
        llama_doc = LlamaDocument(text=text, metadata=metadata)
        
        # Parse into nodes
        nodes = self.node_parser.get_nodes_from_documents([llama_doc])
        
        document_chunks = []
        for i, node in enumerate(nodes):
            chunk_metadata = {
                **metadata,
                "chunk_index": i,
                "node_id": node.node_id,
                "total_chunks": len(nodes)
            }
            
            document_chunks.append(DocumentChunk(
                document_id=metadata.get("document_id", 0),
                chunk_index=i,
                content=node.text,
                metadata=chunk_metadata
            ))
        
        return document_chunks


class EmbeddingService:
    """Handles document embeddings using open-source models"""
    
    def __init__(self):
        self.embeddings = get_embedding_service()
        logger.info(f"Initialized embedding service: {type(self.embeddings).__name__}")
        
        # Get model info for logging
        if hasattr(self.embeddings, 'get_model_info'):
            model_info = self.embeddings.get_model_info()
            logger.info(f"Embedding model info: {model_info}")
    
    async def generate_embeddings(self, texts: List[str]) -> List[List[float]]:
        """Generate embeddings for a list of texts"""
        try:
            logger.info(f"Generating embeddings for {len(texts)} texts")
            embeddings = await self.embeddings.aembed_documents(texts)
            logger.info(f"Successfully generated {len(embeddings)} embeddings")
            return embeddings
        except Exception as e:
            logger.error(f"Error generating embeddings: {str(e)}")
            raise Exception(f"Error generating embeddings: {str(e)}")
    
    async def generate_query_embedding(self, query: str) -> List[float]:
        """Generate embedding for a single query"""
        try:
            embedding = await self.embeddings.aembed_query(query)
            logger.debug(f"Generated query embedding with dimension: {len(embedding)}")
            return embedding
        except Exception as e:
            logger.error(f"Error generating query embedding: {str(e)}")
            raise Exception(f"Error generating query embedding: {str(e)}")
    
    def get_model_info(self) -> Dict[str, Any]:
        """Get information about the embedding model"""
        if hasattr(self.embeddings, 'get_model_info'):
            return self.embeddings.get_model_info()
        return {"model_type": type(self.embeddings).__name__}


class VectorStoreService:
    """Manages vector database operations"""
    
    def __init__(self):
        self.chroma_client = chromadb.PersistentClient(
            path=settings.CHROMA_PERSIST_DIRECTORY,
            settings=ChromaSettings(anonymized_telemetry=False)
        )
        self.collection_name = "knowledge_base"
        self.collection = self._get_or_create_collection()
    
    def _get_or_create_collection(self):
        """Get or create ChromaDB collection"""
        try:
            return self.chroma_client.get_collection(name=self.collection_name)
        except:
            return self.chroma_client.create_collection(
                name=self.collection_name,
                metadata={"hnsw:space": "cosine"}
            )
    
    async def add_documents(self, chunks: List[DocumentChunk], embeddings: List[List[float]]):
        """Add document chunks with embeddings to vector store"""
        try:
            ids = [f"doc_{chunk.document_id}_chunk_{chunk.chunk_index}" for chunk in chunks]
            documents = [chunk.content for chunk in chunks]
            metadatas = [chunk.metadata for chunk in chunks]
            
            self.collection.add(
                ids=ids,
                embeddings=embeddings,
                documents=documents,
                metadatas=metadatas
            )
        except Exception as e:
            raise Exception(f"Error adding documents to vector store: {str(e)}")
    
    async def similarity_search(self, query_embedding: List[float], k: int = 5, 
                               filter_metadata: Dict[str, Any] = None) -> List[Dict[str, Any]]:
        """Search for similar documents"""
        try:
            where_clause = filter_metadata if filter_metadata else None
            
            results = self.collection.query(
                query_embeddings=[query_embedding],
                n_results=k,
                where=where_clause
            )
            
            # Format results
            formatted_results = []
            if results['documents'] and results['documents'][0]:
                for i in range(len(results['documents'][0])):
                    formatted_results.append({
                        'content': results['documents'][0][i],
                        'metadata': results['metadatas'][0][i],
                        'distance': results['distances'][0][i] if results['distances'] else None,
                        'id': results['ids'][0][i]
                    })
            
            return formatted_results
        except Exception as e:
            raise Exception(f"Error performing similarity search: {str(e)}")


class DocumentIngestionPipeline:
    """Main document ingestion pipeline orchestrator"""
    
    def __init__(self):
        self.processor = DocumentProcessor()
        self.chunker = TextChunker()
        self.embedding_service = EmbeddingService()
        self.vector_store = VectorStoreService()
    
    async def ingest_document(self, file_path: str, document_id: int, 
                            metadata: Dict[str, Any] = None) -> Dict[str, Any]:
        """Complete document ingestion pipeline"""
        try:
            start_time = datetime.now()
            
            # Extract file information
            file_info = Path(file_path)
            file_type = file_info.suffix
            file_size = file_info.stat().st_size
            
            if not metadata:
                metadata = {}
            
            metadata.update({
                "document_id": document_id,
                "file_name": file_info.name,
                "file_type": file_type,
                "file_size": file_size,
                "ingestion_timestamp": start_time.isoformat()
            })
            
            # Step 1: Extract text
            text_content = await self.processor.extract_text_from_file(file_path, file_type)
            
            # Step 2: Chunk text
            chunks = self.chunker.chunk_text(text_content, metadata)
            
            # Step 3: Generate embeddings
            chunk_texts = [chunk.content for chunk in chunks]
            embeddings = await self.embedding_service.generate_embeddings(chunk_texts)
            
            # Step 4: Store in vector database
            await self.vector_store.add_documents(chunks, embeddings)
            
            # Calculate processing time
            processing_time = (datetime.now() - start_time).total_seconds()
            
            return {
                "document_id": document_id,
                "chunks_created": len(chunks),
                "total_characters": len(text_content),
                "processing_time": processing_time,
                "status": "success"
            }
            
        except Exception as e:
            return {
                "document_id": document_id,
                "status": "error",
                "error": str(e)
            }
    
    async def search_documents(self, query: str, user_role: str = "employee", 
                             department: str = None, k: int = 5) -> List[Dict[str, Any]]:
        """Search documents with role-based filtering"""
        try:
            # Generate query embedding
            query_embedding = await self.embedding_service.generate_query_embedding(query)
            
            # Build metadata filter based on user role
            filter_metadata = self._build_access_filter(user_role, department)
            
            # Perform similarity search
            results = await self.vector_store.similarity_search(
                query_embedding, k=k, filter_metadata=filter_metadata
            )
            
            return results
            
        except Exception as e:
            raise Exception(f"Error searching documents: {str(e)}")
    
    def _build_access_filter(self, user_role: str, department: str = None) -> Dict[str, Any]:
        """Build access control filter for vector search"""
        filter_conditions = {}
        
        # Role-based access control
        if user_role == "employee":
            filter_conditions["access_level"] = {"$in": ["employee"]}
        elif user_role == "manager":
            filter_conditions["access_level"] = {"$in": ["employee", "manager"]}
        elif user_role == "admin":
            # Admins can access all documents
            pass
        
        # Department-based filtering
        if department and user_role != "admin":
            filter_conditions["department"] = {"$in": [department, None]}
        
        return filter_conditions if filter_conditions else None